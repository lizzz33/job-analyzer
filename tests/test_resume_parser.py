"""Tests for resume_parser — text extraction and LLM parsing."""

import json
from unittest.mock import MagicMock, PropertyMock

import pytest


class TestExtractText:
    def test_extract_from_txt(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text("Hello, I am a developer.", encoding="utf-8")

        from app.services.resume_parser import extract_text_from_file

        assert extract_text_from_file(f) == "Hello, I am a developer."

    def test_extract_from_unsupported_extension(self, tmp_path):
        f = tmp_path / "resume.xlsx"
        f.write_bytes(b"\x00")

        from app.services.resume_parser import extract_text_from_file

        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text_from_file(f)

    def test_extract_from_docx(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Line one")
        doc.add_paragraph("Line two")
        doc.save(str(tmp_path / "resume.docx"))

        from app.services.resume_parser import extract_text_from_file

        text = extract_text_from_file(tmp_path / "resume.docx")
        assert "Line one" in text
        assert "Line two" in text

    def test_extract_from_pdf(self, tmp_path):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        pdf_path = tmp_path / "resume.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        from app.services.resume_parser import extract_text_from_file

        text = extract_text_from_file(pdf_path)
        assert isinstance(text, str)


class TestResumeParser:
    def _make_parser(self, llm_response_content):
        """Build ResumeParser with mocked LLM via property mock."""
        from app.services.resume_parser import ResumeParser

        parser = ResumeParser()
        mock_llm = MagicMock()
        mock_response = MagicMock()
        mock_response.content = llm_response_content
        mock_llm.invoke.return_value = mock_response
        # Patch the property to return our mock
        type(parser).llm = PropertyMock(return_value=mock_llm)
        return parser

    def test_parse_returns_profile_from_llm(self, tmp_path):
        llm_json = json.dumps({
            "name": "Анна Смирнова",
            "position": "Data Scientist",
            "skills": ["Python", "ML", "SQL"],
            "experience_years": "3 г. 6 мес.",
            "education": "МГУ",
            "city": "Москва",
            "summary": "Опытный data scientist.",
        })

        parser = self._make_parser(f'Вот результат:\n{llm_json}\nконец')

        resume_file = tmp_path / "resume.txt"
        resume_file.write_text("Фейковое резюме для теста", encoding="utf-8")

        profile = parser.parse(resume_file)
        assert profile.name == "Анна Смирнова"
        assert "Python" in profile.skills
        assert profile.position == "Data Scientist"
        assert profile.raw_text == "Фейковое резюме для теста"

    def test_parse_empty_file_raises(self, tmp_path):
        from app.services.resume_parser import ResumeParser

        parser = ResumeParser()
        # No need to mock LLM — it should never be called

        empty = tmp_path / "empty.txt"
        empty.write_text("   ", encoding="utf-8")

        with pytest.raises(ValueError, match="Не удалось извлечь текст"):
            parser.parse(empty)

    def test_parse_llm_no_json_returns_fallback(self, tmp_path):
        parser = self._make_parser("Sorry I cannot parse this resume")

        resume_file = tmp_path / "resume.txt"
        resume_file.write_text("Some resume text content here", encoding="utf-8")

        profile = parser.parse(resume_file)
        assert profile.raw_text == "Some resume text content here"
        assert profile.name is None
        assert profile.summary == "Some resume text content here"

    def test_parse_truncates_long_text(self, tmp_path):
        long_text = "A" * 20000
        llm_json = json.dumps({"name": "Test", "skills": []})

        parser = self._make_parser(llm_json)

        resume_file = tmp_path / "long.txt"
        resume_file.write_text(long_text, encoding="utf-8")

        profile = parser.parse(resume_file)
        assert profile.name == "Test"
        assert len(profile.raw_text) == 20000

    def test_parse_llm_invalid_json_uses_fallback(self, tmp_path):
        parser = self._make_parser("This is not JSON at all, no braces")

        resume_file = tmp_path / "resume.txt"
        resume_file.write_text("Valid resume text", encoding="utf-8")

        profile = parser.parse(resume_file)
        assert profile.raw_text == "Valid resume text"
        assert profile.name is None
